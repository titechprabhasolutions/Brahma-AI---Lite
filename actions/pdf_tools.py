"""
pdf_tools.py - Brahma AI PDF support

Creates editable-in-spirit PDF documents from structured content and converts
existing DOCX / text files into readable PDFs without relying on LibreOffice.
"""

from __future__ import annotations

import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

PROJECT_NAME = "Brahma AI - Lite"
DEFAULT_OUTPUT_DIR = Path.home() / "Downloads"


def _sanitize_filename(name: str, default: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._ -]+", "", (name or "").strip())
    safe = re.sub(r"\s+", " ", safe).strip().replace(" ", "_")
    return safe or default


def _resolve_output_path(output_path: str | None, title: str, ext: str, fallback_name: str) -> Path:
    if output_path:
        path = Path(output_path).expanduser()
        if not path.is_absolute():
            head = path.parts[0].lower() if path.parts else ""
            tail = Path(*path.parts[1:]) if len(path.parts) > 1 else Path(path.name)
            if head in {"downloads", "download"}:
                path = Path.home() / "Downloads" / tail
            elif head == "desktop":
                path = Path.home() / "Desktop" / tail
            else:
                path = Path.cwd() / path
        if path.suffix.lower() != ext:
            path = path.with_suffix(ext)
        path.parent.mkdir(parents=True, exist_ok=True)
        return path

    DEFAULT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return DEFAULT_OUTPUT_DIR / f"{_sanitize_filename(title, fallback_name)}{ext}"


def _open_file(path: Path) -> None:
    try:
        if os.name == "nt":
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
    except Exception:
        pass


def _parse_json_arg(value, fallback):
    if value is None:
        return fallback
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return fallback
        try:
            return json.loads(text)
        except Exception:
            return fallback
    return fallback


def _normalize_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return value
    parsed = _parse_json_arg(value, None)
    if isinstance(parsed, list):
        return parsed
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return []
        return [line.strip() for line in re.split(r"\n+", text) if line.strip()]
    return [value]


def _import_pdf():
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            BaseDocTemplate,
            Frame,
            ListFlowable,
            ListItem,
            PageTemplate,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
        return {
            "colors": colors,
            "TA_CENTER": TA_CENTER,
            "TA_JUSTIFY": TA_JUSTIFY,
            "TA_LEFT": TA_LEFT,
            "LETTER": LETTER,
            "ParagraphStyle": ParagraphStyle,
            "getSampleStyleSheet": getSampleStyleSheet,
            "inch": inch,
            "BaseDocTemplate": BaseDocTemplate,
            "Frame": Frame,
            "ListFlowable": ListFlowable,
            "ListItem": ListItem,
            "PageTemplate": PageTemplate,
            "Paragraph": Paragraph,
            "SimpleDocTemplate": SimpleDocTemplate,
            "Spacer": Spacer,
            "Table": Table,
            "TableStyle": TableStyle,
        }
    except Exception as e:
        raise RuntimeError("reportlab is required. Install it with: pip install reportlab") from e


def _open_docx_text(path: Path) -> list[dict]:
    from docx import Document

    doc = Document(path)
    blocks: list[dict] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        kind = "body"
        level = 0
        if style_name.startswith("Heading "):
            kind = "heading"
            try:
                level = int(style_name.split()[-1])
            except Exception:
                level = 1
        blocks.append({"kind": kind, "level": level, "text": text})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                rows.append(row_text)
        if rows:
            blocks.append({"kind": "table", "text": rows})

    return blocks


def _plain_text_blocks(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    blocks = []
    for chunk in [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]:
        blocks.append({"kind": "body", "text": chunk})
    return blocks


def _structured_blocks(parameters: dict) -> list[dict]:
    sections = _parse_json_arg(parameters.get("sections"), None)
    if sections:
        if isinstance(sections, dict):
            sections = [sections]
        blocks = []
        for section in sections:
            if not isinstance(section, dict):
                continue
            heading = (section.get("heading") or section.get("title") or "").strip()
            body = section.get("body") or section.get("content") or ""
            bullets = section.get("bullets") or []
            level = int(section.get("level") or 1)
            if heading:
                blocks.append({"kind": "heading", "level": level, "text": heading})
            if isinstance(body, list):
                for item in body:
                    blocks.append({"kind": "body", "text": str(item)})
            else:
                for para in [p.strip() for p in re.split(r"\n\s*\n", str(body)) if p.strip()]:
                    blocks.append({"kind": "body", "text": para})
            for bullet in _normalize_list(bullets):
                blocks.append({"kind": "bullet", "text": str(bullet)})
        return blocks

    paragraphs = _normalize_list(parameters.get("paragraphs"))
    bullets = _normalize_list(parameters.get("bullets"))
    numbered = _normalize_list(parameters.get("numbered"))
    content = parameters.get("content") or parameters.get("body") or ""

    blocks = []
    if paragraphs:
        for para in paragraphs:
            blocks.append({"kind": "body", "text": str(para)})
    elif content:
        for para in [p.strip() for p in re.split(r"\n\s*\n", str(content)) if p.strip()]:
            blocks.append({"kind": "body", "text": para})

    for bullet in bullets:
        blocks.append({"kind": "bullet", "text": str(bullet)})
    for item in numbered:
        blocks.append({"kind": "numbered", "text": str(item)})

    return blocks


def _render_title_page(story, pdf, title: str, subtitle: str | None, styles):
    story.append(pdf["Spacer"](1, 1.3 * pdf["inch"]))
    story.append(
        pdf["Paragraph"](
            title,
            styles["brahma_doc_title"],
        )
    )
    if subtitle:
        story.append(pdf["Spacer"](1, 0.15 * pdf["inch"]))
        story.append(pdf["Paragraph"](subtitle, styles["brahma_doc_subtitle"]))
    story.append(pdf["Spacer"](1, 0.35 * pdf["inch"]))
    story.append(
        pdf["Paragraph"](
            f"Created by {PROJECT_NAME} on {datetime.now().strftime('%B %d, %Y')}",
            styles["brahma_doc_meta"],
        )
    )
    story.append(pdf["Spacer"](1, 0.5 * pdf["inch"]))


def _pdf_story_from_blocks(blocks: list[dict], pdf, styles):
    story = []
    number_index = 1

    for block in blocks:
        kind = block.get("kind", "body")
        text = str(block.get("text", "")).strip()
        if not text:
            continue
        if kind == "heading":
            number_index = 1
            level = int(block.get("level") or 1)
            story.append(pdf["Paragraph"](text, styles.get(f"brahma_h{min(max(level, 1), 3)}", styles["brahma_h1"])))
            story.append(pdf["Spacer"](1, 0.1 * pdf["inch"]))
        elif kind == "bullet":
            story.append(pdf["Paragraph"](f"- {text}", styles["brahma_body"]))
            story.append(pdf["Spacer"](1, 0.03 * pdf["inch"]))
        elif kind == "numbered":
            story.append(pdf["Paragraph"](f"{number_index}. {text}", styles["brahma_body"]))
            number_index += 1
            story.append(pdf["Spacer"](1, 0.06 * pdf["inch"]))
        elif kind == "table":
            rows = block.get("text") or []
            table_data = [[pdf["Paragraph"](cell, styles["brahma_table_cell"]) for cell in row.split(" | ")] for row in rows]
            if table_data:
                table = pdf["Table"](table_data, repeatRows=0)
                table.setStyle(pdf["TableStyle"]([
                    ("BACKGROUND", (0, 0), (-1, 0), pdf["colors"].HexColor("#0E2230")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), pdf["colors"].white),
                    ("GRID", (0, 0), (-1, -1), 0.5, pdf["colors"].HexColor("#23485E")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(table)
                story.append(pdf["Spacer"](1, 0.15 * pdf["inch"]))
        else:
            story.append(pdf["Paragraph"](text, styles["brahma_body"]))
            story.append(pdf["Spacer"](1, 0.09 * pdf["inch"]))
    return story


def create_pdf(parameters: dict, player=None) -> str:
    pdf = _import_pdf()
    title = (parameters.get("title") or parameters.get("name") or "Document").strip()
    subtitle = (parameters.get("subtitle") or "").strip()
    output_path = _resolve_output_path(parameters.get("output_path"), title, ".pdf", "brahma_ai_output")
    auto_open = parameters.get("auto_open", True)
    action = (parameters.get("action") or "create").lower().strip()
    source_path_str = (parameters.get("file_path") or "").strip()
    source_path = Path(source_path_str) if source_path_str else None

    styles = pdf["getSampleStyleSheet"]()
    styles.add(pdf["ParagraphStyle"](
        name="brahma_doc_title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=pdf["colors"].HexColor("#07131C"),
        alignment=pdf["TA_CENTER"],
        spaceAfter=10,
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_doc_subtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=12,
        leading=15,
        textColor=pdf["colors"].HexColor("#23485E"),
        alignment=pdf["TA_CENTER"],
        spaceAfter=6,
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_doc_meta",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        textColor=pdf["colors"].HexColor("#4D6B7C"),
        alignment=pdf["TA_CENTER"],
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        alignment=pdf["TA_JUSTIFY"],
        spaceAfter=6,
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_h1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        textColor=pdf["colors"].HexColor("#0E2230"),
        spaceBefore=8,
        spaceAfter=6,
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_h2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=15,
        textColor=pdf["colors"].HexColor("#0E2230"),
        spaceBefore=7,
        spaceAfter=5,
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_h3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
        textColor=pdf["colors"].HexColor("#23485E"),
        spaceBefore=6,
        spaceAfter=4,
    ))
    styles.add(pdf["ParagraphStyle"](
        name="brahma_table_cell",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=11.5,
        textColor=pdf["colors"].HexColor("#0E2230"),
    ))

    blocks = []
    if action == "convert" and source_path:
        if not source_path.exists():
            return f"File not found: {source_path}"
        suffix = source_path.suffix.lower()
        if suffix == ".docx":
            blocks = _open_docx_text(source_path)
            if not blocks:
                return "The DOCX appears to be empty."
        elif suffix in {".txt", ".md", ".rst", ".log"}:
            blocks = _plain_text_blocks(source_path)
        else:
            return f"Unsupported source type for PDF conversion: {source_path.suffix}"
    else:
        blocks = _structured_blocks(parameters)

    if not blocks:
        blocks = [{"kind": "body", "text": "Add content here."}]

    doc = pdf["SimpleDocTemplate"](
        str(output_path),
        pagesize=pdf["LETTER"],
        rightMargin=0.8 * pdf["inch"],
        leftMargin=0.8 * pdf["inch"],
        topMargin=0.8 * pdf["inch"],
        bottomMargin=0.8 * pdf["inch"],
        title=title,
        author=parameters.get("author") or "Suryaansh Tiwari",
        subject=parameters.get("subject") or "PDF document",
    )

    story = []
    if title or subtitle:
        _render_title_page(story, pdf, title, subtitle, styles)
        story.append(pdf["Paragraph"]("<hr/>", styles["brahma_body"]))
        story.append(pdf["Spacer"](1, 0.18 * pdf["inch"]))

    if action == "create_letter":
        recipient = (parameters.get("recipient") or parameters.get("to") or "").strip()
        date_value = (parameters.get("date") or datetime.now().strftime("%B %d, %Y")).strip()
        salutation = (parameters.get("salutation") or (f"Dear {recipient}," if recipient else "Dear Sir or Madam,")).strip()
        closing = (parameters.get("closing") or "Sincerely,").strip()
        signature = (parameters.get("signature") or parameters.get("author") or "Suryaansh Tiwari").strip()
        body_text = parameters.get("body") or parameters.get("content") or ""

        story.append(pdf["Paragraph"](date_value, styles["brahma_body"]))
        story.append(pdf["Spacer"](1, 0.08 * pdf["inch"]))
        if recipient:
            story.append(pdf["Paragraph"](recipient, styles["brahma_body"]))
            story.append(pdf["Spacer"](1, 0.06 * pdf["inch"]))
        story.append(pdf["Paragraph"](salutation, styles["brahma_body"]))
        story.append(pdf["Spacer"](1, 0.08 * pdf["inch"]))
        for para in _normalize_list(parameters.get("paragraphs")) or [p.strip() for p in re.split(r"\n\s*\n", str(body_text)) if p.strip()]:
            story.append(pdf["Paragraph"](str(para), styles["brahma_body"]))
            story.append(pdf["Spacer"](1, 0.08 * pdf["inch"]))
        story.append(pdf["Spacer"](1, 0.18 * pdf["inch"]))
        story.append(pdf["Paragraph"](closing, styles["brahma_body"]))
        story.append(pdf["Spacer"](1, 0.3 * pdf["inch"]))
        story.append(pdf["Paragraph"](signature, styles["brahma_body"]))
    else:
        story.extend(_pdf_story_from_blocks(blocks, pdf, styles))

    try:
        doc.build(story)
    except Exception as e:
        return f"PDF creation failed: {e}"

    if auto_open:
        _open_file(output_path)
    return f"PDF created: {output_path}"
