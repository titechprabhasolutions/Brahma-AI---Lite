"""
docx_tools.py - Brahma AI Word / DOCX support

Provides dedicated creation, editing, extraction, summarization, and opening
workflows for editable Word documents.
"""

from __future__ import annotations

import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

PROJECT_NAME = "Brahma AI - Lite"
DEFAULT_OUTPUT_DIR = Path.home() / "Downloads"


def _sanitize_filename(name: str, default: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._ -]+", "", (name or "").strip())
    safe = re.sub(r"\s+", " ", safe).strip().replace(" ", "_")
    return safe or default


def _resolve_output_path(output_path: str | None, title: str, ext: str, fallback_name: str) -> Path:
    if output_path:
        path = Path(output_path).expanduser()
        if not path.is_absolute():
            head = path.parts[0].lower() if path.parts else ""
            tail = Path(*path.parts[1:]) if len(path.parts) > 1 else Path(path.name)
            if head in {"downloads", "download"}:
                path = Path.home() / "Downloads" / tail
            elif head == "desktop":
                path = Path.home() / "Desktop" / tail
            else:
                path = Path.cwd() / path
        if path.suffix.lower() != ext:
            path = path.with_suffix(ext)
        path.parent.mkdir(parents=True, exist_ok=True)
        return path

    DEFAULT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return DEFAULT_OUTPUT_DIR / f"{_sanitize_filename(title, fallback_name)}{ext}"


def _open_file(path: Path) -> None:
    try:
        if os.name == "nt":
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
    except Exception:
        pass


def _get_api_key() -> str:
    config_path = Path(__file__).resolve().parent.parent / "config" / "api_keys.json"
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)["gemini_api_key"]


def _gemini_client():
    import google.generativeai as genai

    genai.configure(api_key=_get_api_key())
    return genai.GenerativeModel("gemini-2.5-flash")


def _import_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        return Document, WD_ALIGN_PARAGRAPH, Inches, Pt, RGBColor
    except Exception as e:
        raise RuntimeError("python-docx is required. Install it with: pip install python-docx") from e


def _parse_json_arg(value, fallback):
    if value is None:
        return fallback
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return fallback
        try:
            return json.loads(text)
        except Exception:
            return fallback
    return fallback


def _normalize_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return value
    parsed = _parse_json_arg(value, None)
    if isinstance(parsed, list):
        return parsed
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return []
        return [line.strip() for line in re.split(r"\n+", text) if line.strip()]
    return [value]


def _set_document_defaults(doc):
    from docx.shared import Inches, Pt

    section = doc.sections[0]
    margin = 1.0
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(margin)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    for style_name, font_name, font_size in [
        ("Title", "Aptos Display", 22),
        ("Subtitle", "Aptos", 14),
        ("Heading 1", "Aptos Display", 16),
        ("Heading 2", "Aptos", 13),
        ("Heading 3", "Aptos", 11),
    ]:
        try:
            style = doc.styles[style_name]
            style.font.name = font_name
            style.font.size = Pt(font_size)
        except Exception:
            pass


def _set_core_props(doc, title: str, author: str | None = None, subject: str | None = None):
    props = doc.core_properties
    props.title = title or PROJECT_NAME
    props.author = author or "Suryaansh Tiwari"
    props.subject = subject or "Word document"
    props.company = PROJECT_NAME
    props.created = datetime.now()


def _paragraph_text(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if style_name.startswith("Heading "):
        try:
            level = int(style_name.split()[-1])
            return f"{'#' * max(1, min(level, 3))} {text}"
        except Exception:
            pass
    return text


def _extract_doc_text(doc) -> str:
    parts = []
    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        if text:
            parts.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                rows.append(row_text)
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _add_paragraph(doc, text: str, style: str | None = None, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def _append_bullets(doc, items):
    for item in items:
        text = item.get("text") if isinstance(item, dict) else str(item)
        if not str(text).strip():
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(str(text).strip())


def _append_numbered(doc, items):
    for item in items:
        text = item.get("text") if isinstance(item, dict) else str(item)
        if not str(text).strip():
            continue
        p = doc.add_paragraph(style="List Number")
        p.add_run(str(text).strip())


def _replace_text_in_runs(doc, replacements: dict[str, str]) -> int:
    changed = 0
    if not replacements:
        return changed
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            original = run.text
            updated = original
            for old, new in replacements.items():
                if old:
                    updated = updated.replace(old, str(new))
            if updated != original:
                run.text = updated
                changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += _replace_text_in_runs(cell, replacements) if hasattr(cell, "paragraphs") else 0
    return changed


def _append_sections(doc, sections):
    for section in sections:
        if not isinstance(section, dict):
            continue
        heading = (section.get("heading") or section.get("title") or "").strip()
        body = section.get("body") or section.get("content") or ""
        bullets = section.get("bullets") or []
        level = int(section.get("level") or 1)

        if heading:
            doc.add_heading(heading, level=max(1, min(level, 3)))

        if isinstance(body, list):
            for item in body:
                _add_paragraph(doc, str(item))
        else:
            body_text = str(body).strip()
            if body_text:
                for para in [p.strip() for p in re.split(r"\n\s*\n", body_text) if p.strip()]:
                    _add_paragraph(doc, para)

        if bullets:
            _append_bullets(doc, _normalize_list(bullets))


def _apply_title_block(doc, title: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = 1
    run = p.add_run(title)
    run.bold = True

    if subtitle:
        sub = doc.add_paragraph()
        sub.style = doc.styles["Subtitle"]
        sub.alignment = 1
        sub.add_run(subtitle)


def _create_letter(doc, params):
    title = (params.get("title") or "Letter").strip()
    recipient = (params.get("recipient") or params.get("to") or "").strip()
    date_value = (params.get("date") or datetime.now().strftime("%B %d, %Y")).strip()
    salutation = (params.get("salutation") or (f"Dear {recipient}," if recipient else "Dear Sir or Madam,")).strip()
    closing = (params.get("closing") or "Sincerely,").strip()
    signature = (params.get("signature") or params.get("author") or "Suryaansh Tiwari").strip()
    body = params.get("body") or params.get("content") or ""
    paragraphs = _normalize_list(params.get("paragraphs")) or [p.strip() for p in re.split(r"\n\s*\n", str(body)) if p.strip()]

    _apply_title_block(doc, title, params.get("subtitle"))
    _add_paragraph(doc, date_value)
    if recipient:
        _add_paragraph(doc, recipient)
    _add_paragraph(doc, salutation)

    for para in paragraphs:
        _add_paragraph(doc, str(para).strip())

    _add_paragraph(doc, closing)
    doc.add_paragraph("")
    _add_paragraph(doc, signature, bold=True)


def _create_report(doc, params):
    title = (params.get("title") or "Report").strip()
    subtitle = (params.get("subtitle") or "").strip()
    content = params.get("content") or params.get("body") or ""
    sections = _parse_json_arg(params.get("sections"), None)
    paragraphs = _normalize_list(params.get("paragraphs"))

    _apply_title_block(doc, title, subtitle)

    if sections:
        if isinstance(sections, dict):
            sections = [sections]
        _append_sections(doc, sections)
    elif paragraphs:
        for para in paragraphs:
            _add_paragraph(doc, str(para))
    elif content:
        for para in [p.strip() for p in re.split(r"\n\s*\n", str(content)) if p.strip()]:
            _add_paragraph(doc, para)
    else:
        _add_paragraph(doc, "Add report content here.")


def _create_generic(doc, params):
    title = (params.get("title") or "Document").strip()
    subtitle = (params.get("subtitle") or "").strip()
    content = params.get("content") or params.get("body") or ""
    sections = _parse_json_arg(params.get("sections"), None)
    paragraphs = _normalize_list(params.get("paragraphs"))
    bullets = _normalize_list(params.get("bullets"))
    numbered = _normalize_list(params.get("numbered"))

    _apply_title_block(doc, title, subtitle)

    if sections:
        if isinstance(sections, dict):
            sections = [sections]
        _append_sections(doc, sections)
        return

    if paragraphs:
        for para in paragraphs:
            _add_paragraph(doc, str(para))
    elif content:
        for para in [p.strip() for p in re.split(r"\n\s*\n", str(content)) if p.strip()]:
            _add_paragraph(doc, para)
    else:
        _add_paragraph(doc, "Add content here.")

    if bullets:
        doc.add_paragraph("")
        _append_bullets(doc, bullets)
    if numbered:
        doc.add_paragraph("")
        _append_numbered(doc, numbered)


def _docx_result_path(source_path: Path | None, action: str, output_path: str | None, title: str) -> Path:
    if output_path:
        fallback = title or (source_path.stem if source_path else "Brahma_AI_Document")
        return _resolve_output_path(output_path, title=fallback, ext=".docx", fallback_name=fallback)
    if source_path:
        return source_path.with_name(f"{source_path.stem}_{action}.docx")
    return _resolve_output_path(None, title=title, ext=".docx", fallback_name="Brahma_AI_Document")


def _load_doc(path: Path):
    Document, _, _, _, _ = _import_docx()
    return Document(path)


def word_document(parameters: dict, player=None, speak=None) -> str:
    params = parameters or {}
    action = (params.get("action") or "create").lower().strip()
    file_path_str = (params.get("file_path") or "").strip()
    output_path_str = (params.get("output_path") or "").strip() or None
    title = (params.get("title") or params.get("subject") or "Brahma AI Document").strip()
    doc_type = (params.get("doc_type") or params.get("template") or "").lower().strip()

    source_path = Path(file_path_str) if file_path_str else None
    if source_path and source_path.suffix.lower() == ".doc":
        return "Legacy .doc files are not supported directly. Please convert the file to .docx first."

    if action in {"open", "open_file", "launch"}:
        if not source_path:
            return "Please provide file_path for the DOCX file to open."
        if not source_path.exists():
            return f"File not found: {source_path}"
        _open_file(source_path)
        return f"Opened {source_path.name} successfully."

    if source_path and not source_path.exists():
        return f"File not found: {source_path}"

    if action in {"read", "inspect"} and source_path:
        try:
            doc = _load_doc(source_path)
            text = _extract_doc_text(doc)
            return text[:8000] if text else "The document appears to be empty."
        except Exception as e:
            return f"Read failed: {e}"

    if action in {"extract_text"} and source_path:
        try:
            doc = _load_doc(source_path)
            text = _extract_doc_text(doc)
            if not text.strip():
                return "The document appears to be empty."
            out = _resolve_output_path(output_path_str, title=source_path.stem, ext=".txt", fallback_name=source_path.stem)
            out = out.with_suffix(".txt")
            out.write_text(text, encoding="utf-8")
            return f"Text extracted. Saved: {out}"
        except Exception as e:
            return f"Extract failed: {e}"

    if action in {"summarize", "analyze"} and source_path:
        try:
            doc = _load_doc(source_path)
            text = _extract_doc_text(doc)
            if not text.strip():
                return "The document appears to be empty."
            model = _gemini_client()
            prompt = (
                "Summarize this Word document concisely and clearly:\n\n"
                if action == "summarize"
                else "Analyze this Word document thoroughly:\n\n"
            )
            response = model.generate_content(prompt + text[:40000])
            result = response.text.strip()
            if len(result) > 600 and params.get("save", True):
                out = _resolve_output_path(output_path_str, title=source_path.stem, ext=".txt", fallback_name=source_path.stem)
                out = out.with_suffix(".txt")
                out.write_text(result, encoding="utf-8")
                return f"{result[:400]}...\n\nFull result saved: {out}"
            return result
        except Exception as e:
            return f"AI analysis failed: {e}"

    Document, WD_ALIGN_PARAGRAPH, Inches, Pt, RGBColor = _import_docx()

    if action in {"append", "add", "edit", "replace_text", "add_heading", "add_bullets", "reformat", "create", "create_letter", "create_report"}:
        if source_path and source_path.exists():
            doc = Document(source_path)
        else:
            doc = Document()
            _set_document_defaults(doc)

        _set_document_defaults(doc)
        _set_core_props(doc, title=title, author=params.get("author"), subject=params.get("subject"))

        created_new = not source_path or not source_path.exists()
        target_path = _docx_result_path(source_path if not created_new else None, action, output_path_str, title)

        if action == "create_letter" or doc_type == "letter":
            _create_letter(doc, params)
        elif action == "create_report" or doc_type == "report":
            _create_report(doc, params)
        elif action == "create":
            _create_generic(doc, params)
        elif action == "reformat":
            if source_path and source_path.exists():
                existing_text = _extract_doc_text(doc)
                doc = Document()
                _set_document_defaults(doc)
                _set_core_props(doc, title=title or source_path.stem, author=params.get("author"), subject=params.get("subject"))
                _apply_title_block(doc, title or source_path.stem, params.get("subtitle"))
                for para in [p.strip() for p in re.split(r"\n\s*\n", existing_text) if p.strip()]:
                    _add_paragraph(doc, para)
            else:
                _create_generic(doc, params)
        else:
            if action in {"replace_text"}:
                replacements = _parse_json_arg(params.get("replacements"), None)
                if not isinstance(replacements, dict):
                    old = params.get("find") or params.get("search")
                    new = params.get("replace") or params.get("replacement") or ""
                    replacements = {str(old): str(new)} if old else {}
                changed = _replace_text_in_runs(doc, replacements or {})
                if changed == 0 and replacements:
                    # fallback for simple documents: replace paragraph text without
                    # forcing a rewrite when there is no run-level match.
                    for paragraph in doc.paragraphs:
                        text = paragraph.text
                        updated = text
                        for old, new in replacements.items():
                            if old:
                                updated = updated.replace(old, str(new))
                        if updated != text:
                            paragraph.text = updated
                            changed += 1
                if not changed:
                    return "No matching text was found to replace."

            if action in {"add_heading"}:
                heading = (params.get("heading") or params.get("text") or params.get("content") or "").strip()
                if not heading:
                    return "Please provide a heading to add."
                level = int(params.get("level") or 1)
                doc.add_heading(heading, level=max(1, min(level, 3)))

            if action in {"add_bullets"}:
                items = _normalize_list(params.get("bullets") or params.get("items") or params.get("content"))
                if not items:
                    return "Please provide bullet items."
                _append_bullets(doc, items)

            if action in {"append", "add", "edit"}:
                heading = (params.get("heading") or "").strip()
                level = int(params.get("level") or 1)
                if heading:
                    doc.add_heading(heading, level=max(1, min(level, 3)))

                body = params.get("content") or params.get("body") or ""
                paragraphs = _normalize_list(params.get("paragraphs"))
                bullets = _normalize_list(params.get("bullets"))
                numbered = _normalize_list(params.get("numbered"))

                if paragraphs:
                    for para in paragraphs:
                        _add_paragraph(doc, str(para))
                elif body:
                    for para in [p.strip() for p in re.split(r"\n\s*\n", str(body)) if p.strip()]:
                        _add_paragraph(doc, para)

                if bullets:
                    _append_bullets(doc, bullets)
                if numbered:
                    _append_numbered(doc, numbered)

        doc.save(target_path)
        if params.get("open_after", True):
            _open_file(target_path)
        return f"Word document created: {target_path}"

    return (
        "Unknown Word action. Try: create, create_letter, create_report, read, summarize, "
        "extract_text, append, replace_text, add_heading, add_bullets, reformat, open"
    )
